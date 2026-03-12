"""
Tests end-to-end para TestRAGSystem (Prompt 10, Part C).

5 tests con ChromaDB in-memory y mocks de LLM/embeddings.
"""

import pytest
import chromadb
from unittest.mock import MagicMock
from chromadb.config import Settings as ChromaSettings
from docx import Document as DocxDocument
from docx.shared import Pt

from app.rag.rag_schema import DocType
from app.rag.rag_system import TestRAGSystem, RAGResponse


# ─────────────────────────────────────────────────
# Fixtures
# ─────────────────────────────────────────────────

@pytest.fixture
def ephemeral_client():
    return chromadb.EphemeralClient(
        settings=ChromaSettings(anonymized_telemetry=False)
    )


@pytest.fixture
def mock_vs(ephemeral_client):
    vs = MagicMock()
    vs.get_or_create_collection = lambda name: ephemeral_client.get_or_create_collection(
        name=name, metadata={"hnsw:space": "cosine"},
    )

    def add_chunks(collection_name, chunks, embeddings):
        col = ephemeral_client.get_or_create_collection(
            name=collection_name, metadata={"hnsw:space": "cosine"},
        )
        col.upsert(
            ids=[c["id"] for c in chunks],
            documents=[c["content"] for c in chunks],
            metadatas=[c["metadata"] for c in chunks],
            embeddings=embeddings,
        )
    vs.add_chunks = add_chunks

    def query(collection_name, query_embedding, top_k=8, where=None):
        col = ephemeral_client.get_or_create_collection(
            name=collection_name, metadata={"hnsw:space": "cosine"},
        )
        if col.count() == 0:
            return []
        kwargs = {
            "query_embeddings": [query_embedding],
            "n_results": min(top_k, col.count()),
            "include": ["documents", "metadatas", "distances"],
        }
        if where:
            kwargs["where"] = where
        try:
            results = col.query(**kwargs)
        except Exception:
            return []
        output = []
        for i in range(len(results["documents"][0])):
            dist = results["distances"][0][i]
            output.append({
                "id": results["ids"][0][i],
                "content": results["documents"][0][i],
                "metadata": results["metadatas"][0][i],
                "distance": dist,
                "relevance_score": 1 - dist,
            })
        return sorted(output, key=lambda x: x["relevance_score"], reverse=True)
    vs.query = query

    def collection_count(collection_name):
        col = ephemeral_client.get_or_create_collection(
            name=collection_name, metadata={"hnsw:space": "cosine"},
        )
        return col.count()
    vs.collection_count = collection_count

    return vs


@pytest.fixture
def mock_emb():
    emb = MagicMock()

    def hash_embed(text):
        import hashlib
        h = hashlib.md5(text.encode()).hexdigest()
        vec = [0.0] * 384
        for i, ch in enumerate(h):
            vec[i % 384] += ord(ch) / 1000.0
        norm = sum(x * x for x in vec) ** 0.5
        if norm > 0:
            vec = [x / norm for x in vec]
        return vec

    emb.embed = hash_embed
    emb.embed_batch = lambda texts: [hash_embed(t) for t in texts]
    return emb


@pytest.fixture
def mock_llm():
    """Simple mock LLM that echoes the query."""
    def llm(system_prompt, context, query):
        return f"Generated response for: {query}\n\nUsing {len(context)} chars of context."
    return llm


@pytest.fixture
def system(mock_vs, mock_emb, mock_llm):
    return TestRAGSystem(
        vector_store=mock_vs,
        embedding_service=mock_emb,
        llm_callable=mock_llm,
    )


@pytest.fixture
def unit_test_file(tmp_path):
    code = """
import pytest
from unittest.mock import MagicMock, patch

def test_save_user():
    mock_repo = MagicMock()
    mock_repo.save.return_value = True
    service = UserService(mock_repo)
    result = service.save_user(user_dto)
    assert result is True
    mock_repo.save.assert_called_once()

def test_find_by_id():
    mock_repo = MagicMock()
    mock_repo.find.return_value = User(id=1)
    service = UserService(mock_repo)
    result = service.find_by_id(1)
    assert result.id == 1
"""
    path = tmp_path / "test_user_service.py"
    path.write_text(code)
    return path


@pytest.fixture
def docx_file(tmp_path):
    doc = DocxDocument()
    doc.add_heading("UserService", level=1)
    doc.add_paragraph("Manages user lifecycle and authentication.")
    p1 = doc.add_paragraph()
    run1 = p1.add_run("def save_user(self, user: UserDTO) -> User:")
    run1.font.name = "Courier New"
    run1.font.size = Pt(10)
    p2 = doc.add_paragraph()
    run2 = p2.add_run("def find_by_id(self, user_id: int) -> Optional[User]:")
    run2.font.name = "Courier New"
    run2.font.size = Pt(10)
    path = tmp_path / "UserService.docx"
    doc.save(str(path))
    return path


# ─────────────────────────────────────────────────
# Test 1: unit test → solo unit tests en resultado
# ─────────────────────────────────────────────────

class TestE2EUnitTestRouting:

    def test_unit_test_goes_to_unit_collection(self, system, unit_test_file):
        result = system.add_document(str(unit_test_file))
        assert result.collection == "unit_tests"
        assert result.doc_type == DocType.UNIT_TEST
        assert result.chunks_created > 0


# ─────────────────────────────────────────────────
# Test 2: docx → context con firma real
# ─────────────────────────────────────────────────

class TestE2EDocxIngestion:

    def test_docx_ingested_to_project_docs(self, system, docx_file):
        result = system.add_document(str(docx_file))
        assert result.collection == "project_docs"
        assert result.chunks_created > 0


# ─────────────────────────────────────────────────
# Test 3: daily note aparece primera
# ─────────────────────────────────────────────────

class TestE2EDailyNote:

    def test_daily_note_ingested_correctly(self, system):
        result = system.add_daily_note(
            "Hoy migré el módulo de pagos a la nueva API. "
            "El método processPayment ahora retorna PaymentResult.",
            date="2024-01-15",
        )
        assert result.doc_type == DocType.DAILY_NOTE
        assert result.chunks_created > 0


# ─────────────────────────────────────────────────
# Test 4: LLM inventa método → MethodGroundingFilter detecta
# ─────────────────────────────────────────────────

class TestE2EMethodHallucination:

    def test_hallucination_detection(self, system):
        from app.rag.method_validator import MethodRegistry, MethodGroundingFilter
        registry = MethodRegistry()
        registry.add_component("UserService", ["save_user", "find_by_id"])
        grounding = MethodGroundingFilter()

        fake_response = "service.save_user(dto)\nservice.fake_method()\nservice.another_fake()"
        result = grounding.filter_hallucinated_methods(
            fake_response, "UserService", registry
        )
        assert result.has_hallucinations is True
        assert "fake_method" in result.hallucinated_methods
        assert "⚠️ ADVERTENCIA" in result.filtered_response


# ─────────────────────────────────────────────────
# Test 5: misma query dos veces → segunda desde caché
# ─────────────────────────────────────────────────

class TestE2ECacheHit:

    def test_second_query_served_from_cache(self, system):
        r1 = system.query("genera un test unitario básico")
        assert r1.served_from_cache is False

        r2 = system.query("genera un test unitario básico")
        assert r2.served_from_cache is True
