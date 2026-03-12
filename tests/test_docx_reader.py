"""
Tests unitarios para docx_reader.py

Crea un .docx fixture en memoria con python-docx y valida
la extracción de secciones, tablas, bloques de código y metadata hints.
"""

import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

from app.rag.docx_reader import DocxReader, DocxContent


# ─────────────────────────────────────────────────
# Fixture: crea un .docx temporal con contenido controlado
# ─────────────────────────────────────────────────

@pytest.fixture
def sample_docx(tmp_path) -> Path:
    """
    Genera un .docx con:
    - Heading 1: "UserService"
    - Párrafo normal de texto
    - Párrafo con fuente Courier New (monospace → código)
    - Heading 2: "Métodos disponibles"
    - Otro párrafo normal
    - Una tabla de 2x3
    """
    doc = Document()

    # ── Heading 1 ──
    doc.add_heading("UserService", level=1)

    # ── Párrafo normal ──
    doc.add_paragraph(
        "Este módulo maneja la lógica de usuarios del sistema. "
        "Incluye autenticación y gestión de perfiles."
    )

    # ── Párrafo con fuente monospace (simula código) ──
    code_para = doc.add_paragraph()
    run = code_para.add_run("def save_user(user: UserDTO) -> User:")
    run.font.name = "Courier New"
    run.font.size = Pt(10)

    code_para2 = doc.add_paragraph()
    run2 = code_para2.add_run("    return self.repository.save(user)")
    run2.font.name = "Courier New"
    run2.font.size = Pt(10)

    # ── Heading 2 ──
    doc.add_heading("Métodos disponibles", level=2)

    # ── Párrafo normal tras segundo heading ──
    doc.add_paragraph(
        "Los métodos principales son save_user, find_by_id y delete_user."
    )

    # ── Tabla 2x3 ──
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Método"
    table.cell(0, 1).text = "Retorno"
    table.cell(1, 0).text = "save_user"
    table.cell(1, 1).text = "User"
    table.cell(2, 0).text = "find_by_id"
    table.cell(2, 1).text = "Optional[User]"

    # Guardar
    file_path = tmp_path / "test_service.docx"
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def reader() -> DocxReader:
    return DocxReader()


# ─────────────────────────────────────────────────
# Tests de read()
# ─────────────────────────────────────────────────

class TestDocxReaderRead:

    def test_returns_docx_content(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert isinstance(result, DocxContent)

    def test_sections_split_by_headings(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        # Debe tener al menos 2 secciones (una por cada heading)
        assert len(result.sections) >= 2

    def test_tables_extracted(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert len(result.tables) == 1
        # La tabla tiene 3 filas y 2 columnas
        assert len(result.tables[0]) == 3
        assert result.tables[0][0] == ["Método", "Retorno"]

    def test_code_blocks_delimited(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert "@@CODE_START@@" in result.raw_text
        assert "@@CODE_END@@" in result.raw_text

    def test_raw_text_contains_content(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert "UserService" in result.raw_text
        assert "save_user" in result.raw_text

    def test_filename_and_path(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert result.filename == "test_service.docx"
        assert result.file_path == str(sample_docx)


# ─────────────────────────────────────────────────
# Tests de metadata_hints
# ─────────────────────────────────────────────────

class TestMetadataHints:

    def test_has_code_blocks_true(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert result.metadata_hints["has_code_blocks"] is True

    def test_has_tables_true(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert result.metadata_hints["has_tables"] is True

    def test_heading_count(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert result.metadata_hints["heading_count"] == 2

    def test_first_heading(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert result.metadata_hints["first_heading"] == "UserService"

    def test_word_count_positive(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert result.metadata_hints["word_count"] > 0

    def test_possible_component_detected(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert result.metadata_hints["possible_component"] == "UserService"

    def test_possible_methods_detected(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        methods = result.metadata_hints["possible_methods"]
        assert "save_user" in methods

    def test_detected_keywords_not_empty(self, reader, sample_docx):
        result = reader.read(str(sample_docx))
        assert len(result.metadata_hints["detected_keywords"]) > 0


# ─────────────────────────────────────────────────
# Tests de read_directory()
# ─────────────────────────────────────────────────

class TestReadDirectory:

    def test_reads_all_docx_files(self, reader, sample_docx, tmp_path):
        # sample_docx ya está en tmp_path
        results = reader.read_directory(str(tmp_path))
        assert len(results) == 1
        assert results[0].filename == "test_service.docx"

    def test_empty_directory(self, reader, tmp_path):
        empty_dir = tmp_path / "empty"
        empty_dir.mkdir()
        results = reader.read_directory(str(empty_dir))
        assert results == []

    def test_nonexistent_directory(self, reader):
        results = reader.read_directory("/tmp/nonexistent_dir_xyz")
        assert results == []


# ─────────────────────────────────────────────────
# Test sin código (control negativo)
# ─────────────────────────────────────────────────

class TestDocxWithoutCode:

    def test_no_code_blocks(self, reader, tmp_path):
        doc = Document()
        doc.add_heading("Introducción", level=1)
        doc.add_paragraph("Este es un párrafo normal sin código.")
        path = tmp_path / "no_code.docx"
        doc.save(str(path))

        result = reader.read(str(path))
        assert result.metadata_hints["has_code_blocks"] is False
        assert "@@CODE_START@@" not in result.raw_text
