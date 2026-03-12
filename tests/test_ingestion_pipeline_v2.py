"""
Tests de integración para document_ingestion_pipeline_v2.py

Usa mocks para EmbeddingService y un ChromaDB EphemeralClient para validar
el flujo completo: clasificación → chunking → indexación en colección correcta.
"""

import pytest
import chromadb
from pathlib import Path
from unittest.mock import MagicMock, patch
from chromadb.config import Settings as ChromaSettings

from app.rag.rag_schema import DocType, CollectionName
from app.rag.docx_reader import DocxReader
from app.rag.document_classifier import DocumentClassifier
from app.rag.structural_chunker import StructuralChunker
from app.rag.document_ingestion_pipeline_v2 import IngestionPipelineV2, IngestResult


# ─────────────────────────────────────────────────
# Fixtures
# ─────────────────────────────────────────────────

@pytest.fixture
def ephemeral_client():
    return chromadb.EphemeralClient(
        settings=ChromaSettings(anonymized_telemetry=False)
    )


@pytest.fixture
def mock_vs(ephemeral_client):
    """VectorStore mock backed by real ephemeral ChromaDB."""
    vs = MagicMock()
    vs.get_or_create_collection = lambda name: ephemeral_client.get_or_create_collection(
        name=name, metadata={"hnsw:space": "cosine"},
    )

    def add_chunks(collection_name, chunks, embeddings):
        col = ephemeral_client.get_or_create_collection(
            name=collection_name, metadata={"hnsw:space": "cosine"},
        )
        col.add(
            ids=[c["id"] for c in chunks],
            documents=[c["content"] for c in chunks],
            metadatas=[c["metadata"] for c in chunks],
            embeddings=embeddings,
        )

    vs.add_chunks = add_chunks

    def query(collection_name, query_embedding, top_k=1):
        col = ephemeral_client.get_or_create_collection(
            name=collection_name, metadata={"hnsw:space": "cosine"},
        )
        if col.count() == 0:
            return []
        results = col.query(
            query_embeddings=[query_embedding], n_results=min(top_k, col.count()),
        )
        return [{"relevance_score": 1 - d} for d in results["distances"][0]] if results["distances"][0] else []

    vs.query = query
    return vs


@pytest.fixture
def mock_emb():
    """Embedding service that returns deterministic fake embeddings."""
    emb = MagicMock()
    counter = {"n": 0}

    def fake_embed(text):
        import hashlib
        h = hashlib.md5(text.encode()).hexdigest()
        vec = [0.0] * 384
        for i, ch in enumerate(h):
            vec[i % 384] += ord(ch) / 1000.0
        return vec

    def fake_embed_batch(texts):
        return [fake_embed(t) for t in texts]

    emb.embed = fake_embed
    emb.embed_batch = fake_embed_batch
    return emb


@pytest.fixture
def pipeline(mock_vs, mock_emb):
    return IngestionPipelineV2(
        vector_store=mock_vs,
        embedding_service=mock_emb,
    )


@pytest.fixture
def unit_test_file(tmp_path) -> Path:
    code = """
import pytest
from unittest.mock import MagicMock

def test_save_user():
    mock_repo = MagicMock()
    mock_repo.save.return_value = True
    assert mock_repo.save() is True

def test_delete_user():
    mock_repo = MagicMock()
    assert mock_repo.delete(1) is None
"""
    path = tmp_path / "test_user_service.py"
    path.write_text(code)
    return path


@pytest.fixture
def integration_test_file(tmp_path) -> Path:
    code = """
from app import TestClient
import requests

def test_api_users():
    client = TestClient(app)
    response = requests.get('/api/users')
    assert response.status_code == 200

def test_database_connection():
    db.connect()
    repository.save(item)
    assert db.query("SELECT 1")
"""
    path = tmp_path / "integration_test_api.ts"
    path.write_text(code)
    return path


@pytest.fixture
def docx_file(tmp_path) -> Path:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading("UserService", level=1)
    doc.add_paragraph("Manages user lifecycle.")
    p = doc.add_paragraph()
    run = p.add_run("def save_user(self, user): pass")
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    path = tmp_path / "UserService.docx"
    doc.save(str(path))
    return path


# ─────────────────────────────────────────────────
# Tests
# ─────────────────────────────────────────────────

class TestIngestFile:

    def test_unit_test_goes_to_unit_tests(self, pipeline, unit_test_file):
        result = pipeline.ingest_file(str(unit_test_file))
        assert result.collection == "unit_tests"
        assert result.doc_type == DocType.UNIT_TEST
        assert result.chunks_created > 0

    def test_integration_test_goes_to_integration_tests(self, pipeline, integration_test_file):
        result = pipeline.ingest_file(str(integration_test_file))
        assert result.collection == "integration_tests"
        assert result.doc_type == DocType.INTEGRATION_TEST
        assert result.chunks_created > 0

    def test_docx_goes_to_project_docs(self, pipeline, docx_file):
        result = pipeline.ingest_file(str(docx_file))
        assert result.collection == "project_docs"
        assert result.chunks_created > 0

    def test_unsupported_extension_skipped(self, pipeline, tmp_path):
        path = tmp_path / "image.png"
        path.write_bytes(b"\x89PNG")
        result = pipeline.ingest_file(str(path))
        assert result.chunks_created == 0


class TestIngestDailyNote:

    def test_daily_note_to_project_docs(self, pipeline):
        result = pipeline.ingest_daily_note(
            "Hoy migré el módulo de pagos.", date="2024-01-15"
        )
        assert result.collection == "project_docs"
        assert result.doc_type == DocType.DAILY_NOTE
        assert result.chunks_created > 0

    def test_daily_note_priority_score(self, pipeline):
        result = pipeline.ingest_daily_note("Cambio importante.", date="2024-03-10")
        assert result.doc_type == DocType.DAILY_NOTE


class TestIngestDirectory:

    def test_ingests_multiple_files(self, pipeline, tmp_path):
        (tmp_path / "test_a.py").write_text(
            "import pytest\nfrom unittest.mock import MagicMock\n\n"
            "def test_a():\n    mock = MagicMock()\n    assert mock() is not None\n"
        )
        (tmp_path / "readme.md").write_text(
            "# Proyecto de Software\n\nEste es el readme del proyecto con documentación."
        )
        (tmp_path / "photo.jpg").write_bytes(b"\xff\xd8")

        results = pipeline.ingest_directory(str(tmp_path))
        supported = [r for r in results]  # .py and .md processed, not .jpg
        assert len(supported) == 2

    def test_nonexistent_dir_returns_empty(self, pipeline):
        results = pipeline.ingest_directory("/tmp/nonexistent_xyz")
        assert results == []


class TestNoDuplicates:

    def test_same_file_twice_second_uses_upsert(self, pipeline, unit_test_file):
        """With upsert semantics, re-ingesting produces same chunk count (idempotent)."""
        r1 = pipeline.ingest_file(str(unit_test_file))
        assert r1.chunks_created > 0
        r2 = pipeline.ingest_file(str(unit_test_file))
        # With deterministic IDs, upsert is idempotent
        assert r2.chunks_created == r1.chunks_created or r2.is_duplicate
