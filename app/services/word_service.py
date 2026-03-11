from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from app.config import get_settings
from app.models.report_model import (
    DailyInput,
    GeneratedReport,
    ReportType,
    TestCaseResult,
    ProjectTask,
)
from app.utils.logger import get_logger

logger = get_logger(__name__)
settings = get_settings()

# ─── Colour palette ──────────────────────────────────────────
PRIMARY = RGBColor(0x1F, 0x49, 0x7D)   # Dark blue
SECONDARY = RGBColor(0x2E, 0x75, 0xB6)  # Medium blue
ACCENT = RGBColor(0x70, 0xAD, 0x47)    # Green
DANGER = RGBColor(0xC0, 0x00, 0x00)    # Red
HEADER_BG = "1F497D"                    # Hex for table headers
PASS_COLOR = "70AD47"
FAIL_COLOR = "C00000"
WARN_COLOR = "FF8C00"

# ─── Column widths (must sum to usable page width) ────────────
# A4 (21 cm) − left margin (3.0 cm) − right margin (2.5 cm) = 15.5 cm
TABLE_TOTAL_W = Cm(15.5)
COL_W = [Cm(3.0), Cm(6.5), Cm(2.5), Cm(3.5)]   # sum = 15.5 cm

# ─── Content style (for table data cells) ─────────────────────
CONTENT_FONT_SIZE = Pt(10)
CONTENT_ITALIC = False


class WordService:
    """
    Builds the complete .docx from a GeneratedReport domain object.
    """

    def __init__(self):
        self._output_dir = Path(settings.output_reports_dir)
        self._output_dir.mkdir(parents=True, exist_ok=True)

    def generate_docx(
        self,
        report: GeneratedReport,
        daily_input: DailyInput,
    ) -> Path:
        """
        Main entry point. Returns path to generated .docx file.
        """
        doc = Document()
        self._configure_document(doc)

        self._build_cover(doc, report, daily_input)
        self._build_toc(doc)
        self._add_page_break(doc)
        self._build_body_sections(doc, report)
        self._build_data_tables(doc, daily_input, report)

        output_path = self._compute_output_path(report)
        doc.save(str(output_path))
        logger.info(f"✅ Document saved: {output_path}")
        return output_path

    # ─────────────────────────────────────────────────
    # Document configuration
    # ─────────────────────────────────────────────────

    def _configure_document(self, doc: Document) -> None:
        """Set page margins and default styles."""
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

        # Set global font to Century Gothic 10pt
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Century Gothic'
        font.size = Pt(10)

        # Header with project identification
        header = section.header
        p = header.paragraphs[0]
        p.text = "auto-report-mcp | Generado automáticamente"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ─────────────────────────────────────────────────
    # Cover page  (metadata table removed)
    # ─────────────────────────────────────────────────

    def _build_cover(
        self, doc: Document, report: GeneratedReport, daily_input: DailyInput
    ) -> None:

        # Report type label
        type_label = {
            ReportType.FUNCTIONAL_TESTS:  "INFORME DE PRUEBAS FUNCIONALES",
            ReportType.INTEGRATION_TESTS: "INFORME DE PRUEBAS DE INTEGRACIÓN",
            ReportType.UNIT_TESTS:        "INFORME DE PRUEBAS UNITARIAS",
            ReportType.PROJECT_PROGRESS:  "INFORME DE AVANCE DE PROYECTO",
        }.get(report.report_type, "INFORME TÉCNICO")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(type_label)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # Project name
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(report.project_name)
        r2.italic = True
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        p2.paragraph_format.space_after = Pt(24)

    # ─────────────────────────────────────────────────
    # Table of Contents
    # ─────────────────────────────────────────────────

    def _build_toc(self, doc: Document) -> None:
        """
        Insert a Word TOC field (updates when the user opens the file and
        accepts the prompt to update fields, or presses Ctrl+A → F9).
        Picks up Heading 2 entries (outlineLevel 1) which are the hidden
        white-text headings placed before each test-case table.
        """
        # TOC title (hidden or removed, since the screenshot doesn't show "Contenido" explicitly, it just shows the table)
        # But we must leave a small spacing if needed, we will just omit the large "Contenido"
        pass

        # TOC field  (\o "2-2" = only Heading 2; \h = hyperlinks; \z = hide tabs in web; \u = use applied outline levels)
        p = doc.add_paragraph()
        run = p.add_run()

        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar_begin)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "2-2" \\h \\z \\u '
        run._r.append(instrText)

        fldChar_sep = OxmlElement('w:fldChar')
        fldChar_sep.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar_sep)

        # Placeholder text (replaced when Word updates the field)
        p_placeholder = doc.add_paragraph()
        p_placeholder.add_run(
            "[Haga clic con el botón derecho aquí y seleccione 'Actualizar campo' para generar el contenido]"
        ).font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        run2 = p.add_run()
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run2._r.append(fldChar_end)

    # ─────────────────────────────────────────────────
    # Hidden heading before each test-case table
    # ─────────────────────────────────────────────────

    def _add_hidden_heading(self, doc: Document, text: str) -> None:
        """
        Adds a Heading 2 paragraph that feeds the TOC. 
        Instead of hidden, it's now visually styled as 9pt black italic.
        """
        p = doc.add_paragraph(style='Heading 2')
        # Remove any existing runs added by the style
        for run in p.runs:
            run.text = ""
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.name = 'Century Gothic'
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)

    # ─────────────────────────────────────────────────
    # Body sections
    # ─────────────────────────────────────────────────

    def _build_body_sections(self, doc: Document, report: GeneratedReport) -> None:
        sorted_sections = sorted(report.sections, key=lambda s: s.section_order)
        for section in sorted_sections:
            doc.add_heading(section.title, level=2)
            for para in section.content.split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

    # ─────────────────────────────────────────────────
    # Data tables
    # ─────────────────────────────────────────────────

    def _build_data_tables(
        self,
        doc: Document,
        daily_input: DailyInput,
        report: GeneratedReport,
    ) -> None:
        rt = report.report_type.value if hasattr(report.report_type, "value") else str(report.report_type)

        _TEST_TYPES = {"functional_tests", "integration_tests", "unit_tests"}
        if rt in _TEST_TYPES and daily_input.test_cases:
            self._build_test_cases_table(
                doc,
                daily_input.test_cases,
                report_type=rt,
            )
        elif rt == "project_progress" and daily_input.tasks:
            self._build_tasks_table(doc, daily_input.tasks)

    def _build_test_cases_table(
        self,
        doc: Document,
        test_cases: list[TestCaseResult],
        report_type: str = "functional_tests",
    ) -> None:
        """
        Unified test-case table builder for all 3 report types.
        Each table is preceded by a hidden Heading 2 (white text, 1pt) that
        feeds the TOC. Cell content uses Century Gothic 9pt italic.
        """
        HEADER_COLOR = "E2EFDA"

        heading_map = {
            "functional_tests":  "PRUEBA(S) FUNCIONAL(ES)",
            "integration_tests": "PRUEBA(S) DE INTEGRACIÓN",
            "unit_tests":        "PRUEBA(S) UNITARIA(S)",
        }
        doc.add_heading(heading_map.get(report_type, "PRUEBA(S)"), level=1)

        extra_rows = {
            "integration_tests": 1,
            "unit_tests": 1,
        }.get(report_type, 0)
        total_rows = 9 + extra_rows

        for i, tc in enumerate(test_cases, 1):
            # ── Hidden heading for TOC ─────────────────────────────────────
            toc_label = f"{i}. {tc.description}" if tc.description else f"{i}. Prueba {tc.test_id}"
            self._add_hidden_heading(doc, toc_label)

            tbl = doc.add_table(rows=total_rows, cols=4)
            tbl.style   = "Table Grid"
            tbl.autofit = False

            self._lock_table_width(tbl, TABLE_TOTAL_W)
            for row in tbl.rows:
                for ci, w in enumerate(COL_W):
                    self._set_cell_width(row.cells[ci], w)

            # ── R1: Header (full-width merge) ──────────────────────────────
            r1 = tbl.rows[0]
            c1 = r1.cells[0]
            c1.merge(r1.cells[3])
            header_label = {
                "functional_tests":  "PRUEBA FUNCIONAL",
                "integration_tests": "PRUEBA DE INTEGRACIÓN",
                "unit_tests":        "PRUEBA UNITARIA (CAJA BLANCA)",
            }.get(report_type, "PRUEBA")
            c1.text = header_label
            self._shade_cell(c1, HEADER_COLOR)
            p1 = c1.paragraphs[0]
            p1.runs[0].bold = True
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # ── R2: Módulo | Número de prueba ──────────────────────────────
            r2 = tbl.rows[1]
            c2_0 = r2.cells[0]
            c2_0.text = "Módulo"
            c2_0.paragraphs[0].runs[0].bold = True
            self._shade_cell(c2_0, "F2F2F2")
            self._vcenter(c2_0)
            r2.cells[1].text = tc.module
            self._apply_content_font(r2.cells[1])
            self._vcenter(r2.cells[1])
            c2_2 = r2.cells[2]
            c2_2.merge(r2.cells[3])
            c2_2.text = f"Número de la prueba {tc.test_id}"
            c2_2.paragraphs[0].runs[0].bold = True
            self._shade_cell(c2_2, "F2F2F2")
            self._vcenter(c2_2)

            # ── R3: Descripción ────────────────────────────────────────────
            r3 = tbl.rows[2]
            c3_0 = r3.cells[0]
            c3_0.text = "Descripción"
            c3_0.paragraphs[0].runs[0].bold = True
            self._shade_cell(c3_0, "F2F2F2")
            self._vcenter(c3_0)
            c3_1 = r3.cells[1]
            c3_1.merge(r3.cells[3])
            c3_1.text = tc.description
            self._apply_content_font(c3_1)

            # ── R4: Preparada por | Fecha ──────────────────────────────────
            r4 = tbl.rows[3]
            c4_0 = r4.cells[0]
            c4_0.text = "Preparada por"
            c4_0.paragraphs[0].runs[0].bold = True
            self._shade_cell(c4_0, "F2F2F2")
            self._vcenter(c4_0)
            r4.cells[1].text = tc.prepared_by or "—"
            self._apply_content_font(r4.cells[1])
            self._vcenter(r4.cells[1])
            c4_2 = r4.cells[2]
            c4_2.text = "Fecha:"
            c4_2.paragraphs[0].runs[0].bold = True
            self._shade_cell(c4_2, "F2F2F2")
            self._vcenter(c4_2)
            r4.cells[3].text = tc.prepare_date or "—"
            self._apply_content_font(r4.cells[3])
            self._vcenter(r4.cells[3])

            # ── R5: Probada por | Fecha ────────────────────────────────────
            r5 = tbl.rows[4]
            c5_0 = r5.cells[0]
            c5_0.text = "Probada por"
            c5_0.paragraphs[0].runs[0].bold = True
            self._shade_cell(c5_0, "F2F2F2")
            self._vcenter(c5_0)
            r5.cells[1].text = tc.tested_by or "—"
            self._apply_content_font(r5.cells[1])
            self._vcenter(r5.cells[1])
            c5_2 = r5.cells[2]
            c5_2.text = "Fecha:"
            c5_2.paragraphs[0].runs[0].bold = True
            self._shade_cell(c5_2, "F2F2F2")
            self._vcenter(c5_2)
            r5.cells[3].text = tc.test_date or "—"
            self._apply_content_font(r5.cells[3])
            self._vcenter(r5.cells[3])

            # ── Extra rows: type-specific metadata ────────────────────────
            next_row = 5

            if report_type == "integration_tests":
                rx = tbl.rows[next_row]
                c_lbl = rx.cells[0]
                c_lbl.text = "Técnica (Caja Negra)"
                c_lbl.paragraphs[0].runs[0].bold = True
                self._shade_cell(c_lbl, "F2F2F2")
                self._vcenter(c_lbl)
                c_val = rx.cells[1]
                c_val.merge(rx.cells[3])
                c_val.text = tc.test_technique or "Partición de equivalencia"
                self._apply_content_font(c_val)
                next_row += 1

            elif report_type == "unit_tests":
                rx = tbl.rows[next_row]
                c_lbl = rx.cells[0]
                c_lbl.text = "Clase / Módulo"
                c_lbl.paragraphs[0].runs[0].bold = True
                self._shade_cell(c_lbl, "F2F2F2")
                self._vcenter(c_lbl)
                c_val = rx.cells[1]
                c_val.merge(rx.cells[3])
                c_val.text = tc.covered_class or "—"
                self._apply_content_font(c_val)
                next_row += 1

            # ── Condiciones de ejecución ──────────────────────────────────
            r6 = tbl.rows[next_row]
            c6 = r6.cells[0]
            c6.merge(r6.cells[3])
            p6_title = c6.paragraphs[0]
            p6_title.clear()
            p6_title.add_run("Condiciones de ejecución").bold = True
            for pc in (list(tc.preconditions or []) or ["—"]):
                pb = c6.add_paragraph(pc, style="List Bullet")
                pb.paragraph_format.space_after = Pt(2)
                self._apply_content_font_paragraph(pb)
            next_row += 1

            # ── Pasos de ejecución ────────────────────────────────────────
            r7 = tbl.rows[next_row]
            c7_0 = r7.cells[0]
            c7_0.text = "Pasos de\nejecución"
            c7_0.paragraphs[0].runs[0].bold = True
            self._shade_cell(c7_0, "F2F2F2")
            c7_1 = r7.cells[1]
            c7_1.merge(r7.cells[3])
            for step in (tc.steps or ["—"]):
                p7 = c7_1.add_paragraph(step, style="List Number")
                p7.paragraph_format.space_after = Pt(2)
                self._apply_content_font_paragraph(p7)
            first_p = c7_1.paragraphs[0]
            if not first_p.text and len(c7_1.paragraphs) > 1:
                first_p._element.getparent().remove(first_p._element)
            next_row += 1

            # ── Resultados esperados ──────────────────────────────────────
            r8 = tbl.rows[next_row]
            c8 = r8.cells[0]
            c8.merge(r8.cells[3])
            p8_title = c8.paragraphs[0]
            p8_title.clear()
            p8_title.add_run("Resultados esperados").bold = True
            for exp in (tc.expected_results or ["—"]):
                pe = c8.add_paragraph(exp, style="List Bullet")
                pe.paragraph_format.space_after = Pt(2)
                self._apply_content_font_paragraph(pe)
            next_row += 1

            # ── Resultados obtenidos ──────────────────────────────────────
            r9 = tbl.rows[next_row]
            c9 = r9.cells[0]
            c9.merge(r9.cells[3])
            p9_title = c9.paragraphs[0]
            p9_title.clear()
            p9_title.add_run("Resultados obtenidos").bold = True

            status_map = {"PASS": "APROBADA", "FAIL": "REPROBADA", "BLOCKED": "BLOQUEADA"}
            local_status = status_map.get(tc.status, tc.status)

            for act in (tc.actual_results or ["—"]):
                pa = c9.add_paragraph(act, style="List Bullet")
                pa.paragraph_format.space_after = Pt(2)
                self._apply_content_font_paragraph(pa)

            p_final = c9.add_paragraph(style="List Bullet")
            p_final.paragraph_format.space_after = Pt(2)
            run_label = p_final.add_run("Resultado de la prueba: ")
            run_label.bold = True
            run_label.font.size = CONTENT_FONT_SIZE
            run_label.font.italic = CONTENT_ITALIC
            run_result = p_final.add_run(local_status)
            run_result.font.size = CONTENT_FONT_SIZE
            run_result.font.italic = CONTENT_ITALIC

            # Evidence image
            if tc.evidence_image_filename:
                img_path = settings.input_images_path / tc.evidence_image_filename
                if img_path.exists():
                    p_img = c9.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        p_img.add_run().add_picture(str(img_path), width=Cm(14.0))
                    except Exception as e:
                        logger.warning(f"No se pudo insertar la imagen {img_path.name}: {e}")

            doc.add_paragraph()

    # ─────────────────────────────────────────────────
    # Content font helpers  (9 pt, italic, Century Gothic)
    # ─────────────────────────────────────────────────

    @staticmethod
    def _apply_content_font(cell) -> None:
        """Apply 9pt italic to all existing runs in all paragraphs of a cell."""
        for para in cell.paragraphs:
            WordService._apply_content_font_paragraph(para)

    @staticmethod
    def _apply_content_font_paragraph(para) -> None:
        """Apply 9pt italic to every run in a paragraph."""
        for run in para.runs:
            if not run.bold:  # leave bold label runs untouched
                run.font.size = CONTENT_FONT_SIZE
                run.font.italic = CONTENT_ITALIC
                run.font.name = 'Century Gothic'

    # ─────────────────────────────────────────────────
    # Tasks table
    # ─────────────────────────────────────────────────

    @staticmethod
    def _label_cell(cell, text: str) -> None:
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        WordService._shade_cell(cell, "F2F2F2")
        WordService._vcenter(cell)

    def _build_tasks_table(
        self, doc: Document, tasks: list[ProjectTask]
    ) -> None:
        doc.add_heading("Detalle de Tareas", level=2)
        headers = ["ID", "Tarea", "Responsable", "Estado", "Avance %", "Sprint"]
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"

        for i, h in enumerate(headers):
            self._style_header_cell(tbl.rows[0].cells[i], text=h)

        for task in tasks:
            row = tbl.add_row()
            row.cells[0].text = task.task_id
            row.cells[1].text = task.title
            row.cells[2].text = task.assignee
            row.cells[3].text = task.status
            row.cells[4].text = f"{task.progress_pct}%"
            row.cells[5].text = task.sprint
            for cell in row.cells:
                self._apply_content_font(cell)

    # ─────────────────────────────────────────────────
    # Helpers
    # ─────────────────────────────────────────────────

    def _compute_output_path(self, report: GeneratedReport) -> Path:
        rt_str = getattr(report.report_type, "value", str(report.report_type))
        filename = (
            f"{report.report_date}_{rt_str}_{report.project_name}"
            .replace(" ", "_")
            .lower()
            + ".docx"
        )
        return self._output_dir / filename

    @staticmethod
    def _add_page_break(doc: Document) -> None:
        from docx.oxml.ns import qn as _qn
        p = doc.add_paragraph()
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(_qn("w:type"), "page")
        run._r.append(br)

    @staticmethod
    def _lock_table_width(tbl, width) -> None:
        from docx.oxml.ns import qn as _qn
        tblPr = tbl._tbl.tblPr
        for old in tblPr.findall(_qn("w:tblW")):
            tblPr.remove(old)
        tblW = OxmlElement("w:tblW")
        tblW.set(_qn("w:w"), str(width.twips))
        tblW.set(_qn("w:type"), "dxa")
        tblPr.append(tblW)
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(_qn("w:type"), "fixed")
        for old in tblPr.findall(_qn("w:tblLayout")):
            tblPr.remove(old)
        tblPr.append(tblLayout)

    @staticmethod
    def _set_cell_width(cell, width) -> None:
        from docx.oxml.ns import qn as _qn
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(_qn("w:tcW")):
            tcPr.remove(old)
        tcW = OxmlElement("w:tcW")
        tcW.set(_qn("w:w"), str(width.twips))
        tcW.set(_qn("w:type"), "dxa")
        tcPr.append(tcW)

    @staticmethod
    def _style_header_cell(cell, text: str | None = None) -> None:
        if text:
            cell.text = text
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
        WordService._shade_cell(cell, HEADER_BG)
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.paragraphs[0].runs[0].bold = True

    @staticmethod
    def _shade_cell(cell, fill_hex: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)

    @staticmethod
    def _vcenter(cell) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement("w:vAlign")
        vAlign.set(qn("w:val"), "center")
        tcPr.append(vAlign)