"""
docx_reader.py — Lector enriquecido de documentos .docx para el RAG v2.

Extrae texto, secciones por heading, tablas, bloques de código (monospace),
y genera metadata hints que alimentan al clasificador y al chunker.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn as _qn

from app.utils.logger import get_logger

logger = get_logger(__name__)

# Fuentes que se consideran monospace / code
_CODE_FONTS = {"courier", "courier new", "consolas", "menlo", "monaco",
               "source code pro", "fira code", "dejavu sans mono",
               "liberation mono", "lucida console"}

_CODE_STYLES = {"code", "código", "macro"}

# Patrones para detectar firmas de función en bloques de código
_METHOD_PATTERNS = [
    re.compile(r"def\s+([a-zA-Z_]\w*)\s*\("),                    # Python
    re.compile(r"function\s+([a-zA-Z_$]\w*)\s*\("),               # JS/TS
    re.compile(r"(?:public|private|protected|static|\s)+\s+\w+\s+([a-zA-Z_]\w*)\s*\("),  # Java/Kotlin/C#
    re.compile(r"(?:async\s+)?([a-zA-Z_$]\w*)\s*\("),             # camelCase genérico
]

# Palabras irrelevantes para detected_keywords
_STOP_WORDS = frozenset({
    "el", "la", "los", "las", "de", "del", "en", "un", "una", "y", "o",
    "que", "es", "se", "con", "para", "por", "al", "su", "sus", "no",
    "the", "a", "an", "and", "or", "of", "in", "to", "is", "it", "for",
    "on", "at", "by", "with", "as", "was", "are", "be", "has", "this",
})


# ── Dataclass de contenido ──────────────────────────────────────

@dataclass
class DocxContent:
    """Resultado de leer un .docx con DocxReader."""
    raw_text: str                             # Texto completo concatenado
    sections: list[str]                       # Dividido por headings
    tables: list[list[list[str]]]             # [tabla][fila][celda]
    metadata_hints: dict                      # Hints para clasificación
    filename: str = ""
    file_path: str = ""


# ── Lector principal ────────────────────────────────────────────

class DocxReader:
    """
    Lee documentos .docx y produce DocxContent con:
    - Secciones divididas por headings
    - Tablas preservadas celda a celda
    - Bloques monospace delimitados con @@CODE_START@@ / @@CODE_END@@
    - Metadata hints para clasificación automática
    """

    def read(self, file_path: str) -> DocxContent:
        """Lee un archivo .docx y retorna DocxContent enriquecido."""
        path = Path(file_path)
        doc = Document(str(path))

        sections: list[str] = []
        current_section_lines: list[str] = []
        tables: list[list[list[str]]] = []
        all_lines: list[str] = []
        in_code_block = False

        body = doc.element.body

        for child in body.iterchildren():
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                para_text = "".join(
                    r.text or "" for r in child.iter(_qn("w:t"))
                ).strip()
                if not para_text:
                    continue

                is_heading = self._is_heading(child)
                is_code = self._is_code_paragraph(child)

                if is_heading:
                    # Cerrar sección anterior
                    if current_section_lines:
                        sections.append("\n".join(current_section_lines))
                        current_section_lines = []
                    # Cerrar bloque de código si estaba abierto
                    if in_code_block:
                        all_lines.append("@@CODE_END@@")
                        current_section_lines.append("@@CODE_END@@")
                        in_code_block = False
                    # El heading empieza nueva sección
                    current_section_lines.append(para_text)
                    all_lines.append(para_text)

                elif is_code:
                    if not in_code_block:
                        all_lines.append("@@CODE_START@@")
                        current_section_lines.append("@@CODE_START@@")
                        in_code_block = True
                    all_lines.append(para_text)
                    current_section_lines.append(para_text)

                else:
                    if in_code_block:
                        all_lines.append("@@CODE_END@@")
                        current_section_lines.append("@@CODE_END@@")
                        in_code_block = False
                    all_lines.append(para_text)
                    current_section_lines.append(para_text)

            elif tag == "tbl":
                if in_code_block:
                    all_lines.append("@@CODE_END@@")
                    current_section_lines.append("@@CODE_END@@")
                    in_code_block = False

                table_data = self._extract_table(child)
                if table_data:
                    tables.append(table_data)
                    # También agregar como texto a la sección actual
                    for row in table_data:
                        row_text = " | ".join(row)
                        all_lines.append(row_text)
                        current_section_lines.append(row_text)

        # Cerrar última sección y posible bloque de código
        if in_code_block:
            all_lines.append("@@CODE_END@@")
            current_section_lines.append("@@CODE_END@@")
        if current_section_lines:
            sections.append("\n".join(current_section_lines))

        raw_text = "\n".join(all_lines)
        metadata_hints = self._extract_metadata_hints(doc, raw_text, sections)

        return DocxContent(
            raw_text=raw_text,
            sections=sections,
            tables=tables,
            metadata_hints=metadata_hints,
            filename=path.name,
            file_path=str(path),
        )

    def read_directory(self, dir_path: str) -> list[DocxContent]:
        """Lee todos los .docx de un directorio. Loggea archivos que fallen."""
        directory = Path(dir_path)
        results: list[DocxContent] = []

        if not directory.exists():
            logger.warning(f"Directory not found: {dir_path}")
            return results

        for docx_path in sorted(directory.glob("*.docx")):
            try:
                content = self.read(str(docx_path))
                results.append(content)
                logger.debug(f"Read: {docx_path.name}")
            except Exception as e:
                logger.error(
                    f"Failed to read {docx_path.name}: {e}",
                    exc_info=True,
                )

        logger.info(
            f"Read {len(results)} .docx file(s) from {dir_path}"
        )
        return results

    # ─────────────────────────────────────────────────
    # Metadata hints
    # ─────────────────────────────────────────────────

    @staticmethod
    def _extract_metadata_hints(
        doc: Document,
        raw_text: str,
        sections: list[str],
    ) -> dict:
        """Genera hints para alimentar al clasificador automático."""
        headings: list[str] = []
        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower() if para.style else ""
            if "heading" in style_name or "título" in style_name or "encabezado" in style_name:
                headings.append(para.text.strip())

        has_code = "@@CODE_START@@" in raw_text
        has_tables = bool(doc.tables)
        word_count = len(raw_text.split())

        # Primeras 20 palabras significativas únicas
        words = raw_text.split()
        seen: set[str] = set()
        keywords: list[str] = []
        for w in words:
            clean = re.sub(r"[^\w]", "", w).lower()
            if clean and len(clean) > 2 and clean not in _STOP_WORDS and clean not in seen:
                seen.add(clean)
                keywords.append(clean)
            if len(keywords) >= 20:
                break

        # Buscar componente en primeros 3 headings
        possible_component = ""
        component_patterns = [
            re.compile(r"(?:Clase|Módulo|Module|Class|Service|Repository|Controller)[:\s]+(\w+)", re.IGNORECASE),
            re.compile(r"(\w+(?:Service|Repository|Controller|Manager|Handler|Factory|Provider|Adapter))", re.IGNORECASE),
        ]
        for heading in headings[:3]:
            for pattern in component_patterns:
                match = pattern.search(heading)
                if match:
                    possible_component = match.group(1)
                    break
            if possible_component:
                break

        # Buscar métodos en bloques de código
        possible_methods: list[str] = []
        code_blocks = re.findall(
            r"@@CODE_START@@(.*?)@@CODE_END@@", raw_text, re.DOTALL
        )
        for block in code_blocks:
            for pattern in _METHOD_PATTERNS:
                for match in pattern.finditer(block):
                    method = match.group(1)
                    # Filtrar keywords comunes que no son métodos reales
                    if method not in {"if", "for", "while", "return", "class",
                                      "def", "function", "async", "await",
                                      "import", "from", "var", "let", "const"}:
                        if method not in possible_methods:
                            possible_methods.append(method)

        return {
            "has_code_blocks": has_code,
            "has_tables": has_tables,
            "heading_count": len(headings),
            "first_heading": headings[0] if headings else "",
            "word_count": word_count,
            "detected_keywords": keywords,
            "possible_component": possible_component,
            "possible_methods": possible_methods,
        }

    # ─────────────────────────────────────────────────
    # Helpers internos
    # ─────────────────────────────────────────────────

    @staticmethod
    def _is_heading(paragraph_element) -> bool:
        """Detecta si un párrafo XML es un heading."""
        pPr = paragraph_element.find(_qn("w:pPr"))
        if pPr is None:
            return False
        pStyle = pPr.find(_qn("w:pStyle"))
        if pStyle is not None:
            style_val = (pStyle.get(_qn("w:val")) or "").lower()
            if any(k in style_val for k in ("heading", "título", "encabezado")):
                return True
        return False

    @staticmethod
    def _is_code_paragraph(paragraph_element) -> bool:
        """Detecta si un párrafo tiene fuente monospace o estilo de código."""
        # Verificar estilo del párrafo
        pPr = paragraph_element.find(_qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(_qn("w:pStyle"))
            if pStyle is not None:
                style_val = (pStyle.get(_qn("w:val")) or "").lower()
                if any(k in style_val for k in _CODE_STYLES):
                    return True

        # Verificar fuente de los runs
        for run in paragraph_element.iter(_qn("w:r")):
            rPr = run.find(_qn("w:rPr"))
            if rPr is not None:
                for font_tag in [_qn("w:rFonts")]:
                    rFonts = rPr.find(font_tag)
                    if rFonts is not None:
                        for attr in ["w:ascii", "w:hAnsi", "w:cs"]:
                            font_name = (rFonts.get(_qn(attr)) or "").lower()
                            if font_name in _CODE_FONTS:
                                return True
        return False

    @staticmethod
    def _extract_table(table_element) -> list[list[str]]:
        """Extrae una tabla del XML como matriz [fila][celda]."""
        rows: list[list[str]] = []
        for row in table_element.iter(_qn("w:tr")):
            cells: list[str] = []
            for cell in row.iter(_qn("w:tc")):
                cell_text = "".join(
                    t.text or "" for t in cell.iter(_qn("w:t"))
                ).strip()
                cells.append(cell_text)
            if cells:
                rows.append(cells)
        return rows
