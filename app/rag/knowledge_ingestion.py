from __future__ import annotations

import hashlib
import json
from datetime import datetime, timezone
from pathlib import Path

from app.config import get_settings
from app.rag.embedding_service import EmbeddingService
from app.rag.vector_store import VectorStore
from app.utils.logger import get_logger

logger = get_logger(__name__)
settings = get_settings()

_REGISTRY_PATH = Path("./data/knowledge_processed.json")
_CHUNK_WORDS   = 150   # smaller chunks keep method names + context together
_CHUNK_OVERLAP = 30


class KnowledgeIngestionPipeline:
    """
    Ingests two types of sources into the `project_knowledge` ChromaDB collection:
    - Free-text notes written by the user  → ingest_text_note()
    - .docx context reports in context_reports/ → ingest_all_context_reports()

    Already-processed files are tracked via a local JSON registry so they are
    never ingested twice.

    Metadata strategy (Recency Weighting):
    - Every chunk carries a `timestamp` (ISO-8601) and a `type` field.
      type="note"           → manual user note  (highest priority at retrieval time)
      type="context_report" → auto-generated .docx report
    - Each chunk is prefixed with a dated tag so the LLM reads the recency signal
      inline before extracting information.
    """

    def __init__(self) -> None:
        self._vs         = VectorStore()
        self._emb        = EmbeddingService()
        self._collection = settings.chroma_collection_knowledge

    # ─────────────────────────────────────────────────
    # Public API
    # ─────────────────────────────────────────────────

    def ingest_text_note(self, note: str) -> int:
        """Proxies note ingestion to TestRAGSystem."""
        from app.rag.rag_system import TestRAGSystem
        sys = TestRAGSystem()
        res = sys.add_daily_note(note)
        return res.chunks_created

    def find_similar_notes(self, text: str, threshold: float = 0.4) -> list[dict]:
        """Proxy to search daily notes via VectorStore."""
        if not text.strip():
            return []
        emb = self._emb.embed_batch([text])[0]
        results = self._vs.query(
            collection_name="project_docs",
            query_embedding=emb,
            top_k=5,
            where={"doc_type": "DAILY_NOTE"},
        )
        return [r for r in results if r.get("relevance_score", 0) >= threshold]

    def delete_notes(self, note_ids: list[str]) -> None:
        if not note_ids:
            return
        self._vs.delete_chunks("project_docs", ids=note_ids)
        logger.info(f"Deleted {len(note_ids)} notes from project_docs")

    def ingest_all_context_reports(self, force: bool = False) -> dict[str, int]:
        """Proxies directory ingestion to TestRAGSystem."""
        context_dir = Path(settings.context_reports_dir)
        if not context_dir.exists():
            return {}
        
        from app.rag.rag_system import TestRAGSystem
        sys = TestRAGSystem()
        
        # Ingest directory uses IngestionPipelineV2 which natively skips real duplicates
        results = sys.add_directory(str(context_dir), recursive=False)
        return {r.source: r.chunks_created for r in results if r.chunks_created > 0}


    def force_reingest_file(self, docx_path: Path) -> int:
        """Proxies force reingestion by just calling add_document."""
        from app.rag.rag_system import TestRAGSystem
        sys = TestRAGSystem()
        res = sys.add_document(str(docx_path))
        return res.chunks_created

    # ─────────────────────────────────────────────────
    # Text extraction
    # ─────────────────────────────────────────────────

    @staticmethod
    def _extract_docx_text(path: Path) -> str:
        try:
            from docx import Document
            doc        = Document(str(path))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Failed to extract text from {path.name}: {e}")
            return ""

    # ─────────────────────────────────────────────────
    # Chunking
    # ─────────────────────────────────────────────────

    @staticmethod
    def _chunk_text(
        text: str,
        chunk_words: int = _CHUNK_WORDS,
        overlap_words: int = _CHUNK_OVERLAP,
    ) -> list[str]:
        words  = text.split()
        chunks = []
        i      = 0
        while i < len(words):
            chunk = " ".join(words[i : i + chunk_words])
            if chunk.strip():
                chunks.append(chunk)
            i += chunk_words - overlap_words
        return chunks
